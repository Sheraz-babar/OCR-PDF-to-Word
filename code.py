from pdf2image import convert_from_path
from PIL import Image, ImageOps
import pytesseract
from docx import Document
import time
import os

pdf_file = "Your_file_path_here.pdf"
output_docx = "Your_output_file.docx"
tesseract_path = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
# change the above path to your tesseract executable if it's different

pytesseract.pytesseract.tesseract_cmd = tesseract_path

# Start timer
start_time = time.time()

print("\n-------------------------------\n🚀 WELCOME TO THE OCR ENGINE!\n-------------------------------\n")
print(f"📄 Input File: {pdf_file}")

# Check if file exists
if not os.path.exists(pdf_file):
    print("❌ Error: PDF file not found!")
    exit()

print("📥 Converting PDF to images...")

# Convert PDF to images
pages = convert_from_path(pdf_file, dpi=300)
total_pages = len(pages)

print(f"✅ Conversion Completed!")
print(f"📊 Total Pages Found: {total_pages}")
print(f"📁 File Processed: {pdf_file}\n")

# Create Word document
doc = Document()

print("🧠 Starting OCR processing...\n")

for i, page in enumerate(pages):
    print(f"Processing page {i+1}/{total_pages}...")

    # Preprocess image for OCR
    gray = ImageOps.grayscale(page)
    bw = gray.point(lambda x: 0 if x < 150 else 255, '1')

    # OCR
    text = pytesseract.image_to_string(bw, config='--psm 6')

    # Add to Word doc
    doc.add_paragraph(f"--- Page {i+1} ---")
    doc.add_paragraph(text)
    doc.add_page_break()

    print(f"Page {i+1} done.\n")

# Save Word file
print("💾 Saving Word document...")
doc.save(output_docx)

# End timer
end_time = time.time()
total_time = end_time - start_time

minutes = int(total_time // 60)
seconds = int(total_time % 60)

print("\n🎉 All pages processed successfully!")
print(f"📄 Output File: {output_docx}")
print(f"\n⏱️ Total time taken: {minutes} minutes {seconds} seconds")